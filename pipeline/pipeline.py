"""
pipeline.py — Computational replication of Bressman & Stack,
"Regulatory Settlement, Stare Decisis, and Loper Bright" (100 NYU L. Rev. 1799).

READ BRIEFING.md FIRST. This file is a runnable skeleton; a coding agent
should flesh out the TODOs, verify the API shapes, and add error handling.

Overview:
  Stage 1a  pull_opinions()              CourtListener  -> data/opinions/*.json
  Stage 1b  classify_opinions()          Anthropic API  -> data/classified_opinions.jsonl
  Stage 1c  validate_named_cases()       sanity check vs. paper
  Stage 2a  pull_rule_histories()        Federal Register API -> data/amendments/*.json
  Stage 2b  classify_amendments()        Anthropic API  -> data/classified_amendments.jsonl
  Stage 3   package()                    build cases.xlsx, amendments.xlsx, chart.png, memo.docx

Run stage by stage:
  python pipeline.py pull
  python pipeline.py classify
  python pipeline.py validate
  python pipeline.py history
  python pipeline.py code-amendments
  python pipeline.py package
  python pipeline.py smoke   # limits everything to 2016 for validation
"""

from __future__ import annotations

import json
import os
import re
import sys
import time
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any, Iterator

import requests

# ------------------------------------------------------------------ config

DATA = Path("data")
OUT = Path("outputs")
DATA.mkdir(exist_ok=True)
OUT.mkdir(exist_ok=True)
(DATA / "opinions").mkdir(exist_ok=True)
(DATA / "amendments").mkdir(exist_ok=True)

COURTLISTENER_BASE = "https://www.courtlistener.com/api/rest/v4"
FEDERAL_REGISTER_BASE = "https://www.federalregister.gov/api/v1"

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY")
COURTLISTENER_TOKEN = os.environ.get("COURTLISTENER_TOKEN")  # optional but recommended

DATE_FROM = "2004-01-01"
DATE_TO = "2024-01-02"
SMOKE_FROM = "2015-01-01"
SMOKE_TO = "2016-12-31"

MODEL = "claude-sonnet-4-6"  # adjust to whatever is current/available

# ------------------------------------------------------------------ stage 1a: pull opinions

def cl_headers():
    h = {"User-Agent": "regulatory-replication/0.1 (academic use)"}
    if COURTLISTENER_TOKEN:
        h["Authorization"] = f"Token {COURTLISTENER_TOKEN}"
    return h


def pull_opinions(date_from: str = DATE_FROM, date_to: str = DATE_TO) -> int:
    """Pull every D.C. Circuit opinion in [date_from, date_to] whose text contains 'Chevron'.

    CourtListener v4 search returns cluster hits; each hit may embed an `opinions`
    list with individual opinion IDs. Full plain text is typically NOT inline —
    we always follow up with a per-opinion fetch to get the body."""
    url = f"{COURTLISTENER_BASE}/search/"
    params = {
        "q": "Chevron",
        "court": "cadc",
        "type": "o",
        "filed_after": date_from,
        "filed_before": date_to,
        "order_by": "dateFiled asc",
        "page_size": 100,
    }
    count = 0
    while url:
        r = requests.get(url, params=params if "?" not in url else None, headers=cl_headers(), timeout=30)
        r.raise_for_status()
        data = r.json()
        for hit in data.get("results", []):
            cluster_id = hit.get("cluster_id") or hit.get("id")
            sub_opinions = hit.get("opinions") or []
            if not sub_opinions:
                sub_opinions = [{"id": None}]
            for sub in sub_opinions:
                opinion_id = sub.get("id")
                key = opinion_id or cluster_id
                if not key:
                    continue
                dst = DATA / "opinions" / f"{key}.json"
                if dst.exists():
                    count += 1
                    continue
                text = ""
                if opinion_id:
                    text = fetch_opinion_text(opinion_id)
                if not text and cluster_id:
                    text = fetch_cluster_text(cluster_id)
                if not text:
                    text = sub.get("snippet") or hit.get("plain_text") or ""
                payload = {
                    "id": key,
                    "cluster_id": cluster_id,
                    "opinion_id": opinion_id,
                    "case_name": hit.get("caseName"),
                    "date_filed": hit.get("dateFiled"),
                    "court": hit.get("court"),
                    "docket": hit.get("docketNumber"),
                    "citations": hit.get("citation"),
                    "absolute_url": hit.get("absolute_url"),
                    "plain_text": text,
                }
                dst.write_text(json.dumps(payload, indent=2))
                count += 1
        url = data.get("next")
        params = None  # next URL already has params
        time.sleep(1.0)  # be polite
    print(f"pulled {count} opinions into {DATA/'opinions'}")
    return count


def _opinion_text_from_json(o: dict) -> str:
    """Extract the best text representation from a CourtListener opinion object.

    The CourtListener docs explicitly recommend `html_with_citations` over
    `plain_text`: `plain_text` is only populated when the source was a PDF or
    Word doc, while `html_with_citations` is the unified rendered field that's
    always populated. We fall through the documented HTML variants and only
    use `plain_text` as a last resort."""
    for html_key in ("html_with_citations", "html", "html_lawbox", "html_columbia", "html_anon_2020", "xml_harvard"):
        html = o.get(html_key) or ""
        if html:
            return re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", html)).strip()
    return o.get("plain_text") or ""


# Spacing between opinion/cluster fetches. CourtListener's authenticated rate
# limit is 5,000/hr (~0.72s min); 0.75s keeps us safely under. The per-search-page
# sleep in pull_opinions is separate from this.
_FETCH_SPACING_S = 0.75


def fetch_opinion_text(opinion_id: int) -> str:
    """Fetch full text for an opinion, throttled to respect CourtListener's rate limit."""
    r = requests.get(f"{COURTLISTENER_BASE}/opinions/{opinion_id}/", headers=cl_headers(), timeout=30)
    time.sleep(_FETCH_SPACING_S)
    if r.status_code != 200:
        return ""
    return _opinion_text_from_json(r.json())


def fetch_cluster_text(cluster_id: int) -> str:
    """Fallback path: resolve opinions through the cluster endpoint when the search
    hit did not embed an opinion id we can fetch directly."""
    r = requests.get(f"{COURTLISTENER_BASE}/clusters/{cluster_id}/", headers=cl_headers(), timeout=30)
    time.sleep(_FETCH_SPACING_S)
    if r.status_code != 200:
        return ""
    cluster = r.json()
    for op_url in cluster.get("sub_opinions", []) or []:
        try:
            rr = requests.get(op_url, headers=cl_headers(), timeout=30)
            time.sleep(_FETCH_SPACING_S)
            if rr.status_code != 200:
                continue
            text = _opinion_text_from_json(rr.json())
            if text:
                return text
        except requests.RequestException:
            continue
    return ""


# ------------------------------------------------------------------ stage 1b: classify opinions

CLASSIFIER_PROMPT = Path("prompts").joinpath("classifier_opinion.md")  # created by prompts.md

_BATCH_CHUNK = 50_000  # well under the 100K/256MB Batches API caps
_BATCH_POLL_S = 60


def _submit_and_wait_for_batch(client, batch_requests, label: str):
    """Submit a batch and poll until `processing_status == 'ended'`."""
    batch = client.messages.batches.create(requests=batch_requests)
    print(f"[{label}] batch {batch.id} submitted ({len(batch_requests)} requests)")
    while True:
        b = client.messages.batches.retrieve(batch.id)
        if b.processing_status == "ended":
            break
        c = b.request_counts
        print(f"[{label}]   status={b.processing_status} proc={c.processing} ok={c.succeeded} err={c.errored} exp={c.expired}")
        time.sleep(_BATCH_POLL_S)
    c = b.request_counts
    print(f"[{label}] complete: ok={c.succeeded} err={c.errored} exp={c.expired}")
    return b


def classify_opinions() -> None:
    """Classify every cached opinion via the Batches API (50% cheaper, async).

    Typical wall time: ~1 hour regardless of batch size, vs. multi-hour sequential.
    Results arrive together at batch end; append-only write to the same jsonl
    keeps resume semantics — a crash or interrupt just re-runs from what's on disk."""
    import anthropic  # pip install anthropic
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    prompt_template = CLASSIFIER_PROMPT.read_text() if CLASSIFIER_PROMPT.exists() else OPINION_CLASSIFIER_FALLBACK
    static_prefix, varying_template = _split_for_cache(prompt_template)

    out_path = DATA / "classified_opinions.jsonl"
    seen: set = set()
    if out_path.exists():
        for line in out_path.read_text().splitlines():
            try:
                seen.add(json.loads(line)["id"])
            except Exception:
                pass

    batch_requests: list = []
    id_to_op: dict = {}
    for opinion_file in sorted((DATA / "opinions").glob("*.json")):
        op = json.loads(opinion_file.read_text())
        if op["id"] in seen or not op.get("plain_text"):
            continue
        text = op["plain_text"][:40_000]
        varying = (
            varying_template
            .replace("{{CASE_NAME}}", op.get("case_name") or "")
            .replace("{{OPINION_TEXT}}", text)
        )
        custom_id = f"op-{op['id']}"
        id_to_op[custom_id] = op
        batch_requests.append({
            "custom_id": custom_id,
            "params": {
                "model": MODEL,
                "max_tokens": 800,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": static_prefix, "cache_control": {"type": "ephemeral"}},
                        {"type": "text", "text": varying},
                    ],
                }],
            },
        })

    if not batch_requests:
        print("no new opinions to classify")
        return
    print(f"queueing {len(batch_requests)} opinions for classification via Batches API")

    with out_path.open("a") as out:
        for i in range(0, len(batch_requests), _BATCH_CHUNK):
            chunk = batch_requests[i:i + _BATCH_CHUNK]
            label = f"opinions {i+1}-{i+len(chunk)}"
            batch = _submit_and_wait_for_batch(client, chunk, label)
            for result in client.messages.batches.results(batch.id):
                op = id_to_op.get(result.custom_id)
                if not op:
                    continue
                if result.result.type != "succeeded":
                    err = getattr(getattr(result.result, "error", None), "type", result.result.type)
                    print(f"  {result.custom_id} failed: {err}")
                    continue
                msg = result.result.message
                parsed = parse_json_block(msg.content[0].text)
                record = {
                    "id": op["id"],
                    "cluster_id": op.get("cluster_id"),
                    "case_name": op["case_name"],
                    "date_filed": op["date_filed"],
                    **parsed,
                }
                out.write(json.dumps(record) + "\n")
                out.flush()


def parse_json_block(text: str) -> dict:
    """Extract and parse the first balanced {...} JSON object from LLM output.

    Handles markdown code fences and trailing prose. Scans with brace-depth
    counting and string-aware escaping so it doesn't trip on `}` inside string
    values or on multiple JSON blocks separated by commentary."""
    if not text:
        return {"error": "empty_response", "raw": ""}
    fence = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.DOTALL)
    if fence:
        try:
            return json.loads(fence.group(1))
        except json.JSONDecodeError:
            pass
    start = text.find("{")
    while start != -1:
        depth = 0
        in_string = False
        escape = False
        for i in range(start, len(text)):
            ch = text[i]
            if escape:
                escape = False
                continue
            if ch == "\\":
                escape = True
                continue
            if ch == '"':
                in_string = not in_string
                continue
            if in_string:
                continue
            if ch == "{":
                depth += 1
            elif ch == "}":
                depth -= 1
                if depth == 0:
                    try:
                        return json.loads(text[start:i + 1])
                    except json.JSONDecodeError:
                        break
        start = text.find("{", start + 1)
    return {"error": "no_parseable_json", "raw": text[:500]}


def _split_for_cache(template: str) -> tuple[str, str]:
    """Split a prompt template at the first `{{placeholder}}` marker.

    The leading static portion is eligible for prompt caching across requests;
    the varying tail is filled per-item. Sonnet 4.6's cache minimum is 2048
    tokens — prefixes shorter than that silently won't cache (no error; usage
    just reports cache_read_input_tokens=0). Structuring this way means caching
    kicks in automatically if the template grows (e.g., adding few-shot
    examples per the iteration notes in prompts.md)."""
    idx = template.find("{{")
    if idx == -1:
        return template, ""
    return template[:idx], template[idx:]


OPINION_CLASSIFIER_FALLBACK = """You are classifying a D.C. Circuit opinion for a legal empirical study.

We want to identify opinions in which the D.C. Circuit UPHELD an agency's
interpretation of an ambiguous statute at CHEVRON STEP TWO, where that
interpretation was adopted through a NOTICE-AND-COMMENT RULEMAKING.

Read the opinion below and return a single JSON object with these fields:

  is_chevron_step_two_affirmance: boolean  // did the court apply Chevron step two and uphold the agency?
  is_notice_and_comment_rule: boolean      // was the interpretation in a notice-and-comment regulation (not adjudication)?
  confidence: number                       // 0.0 to 1.0 overall confidence in the above
  cfr_citation: string                     // best CFR citation for the rule, e.g. "34 C.F.R. pts. 600, 668" (or "")
  federal_register_citation: string        // best Federal Register citation, e.g. "79 Fed. Reg. 64890" (or "")
  agency: string                           // issuing agency
  rule_short_name: string                  // short plain-English name of the rule
  reasoning: string                        // 2-3 sentence explanation

CASE: {{CASE_NAME}}

OPINION:
{{OPINION_TEXT}}

Respond with ONLY the JSON object, no prose before or after.
"""


# ------------------------------------------------------------------ stage 1c: validate

# Named-case matchers. CourtListener's stored case names don't always match the
# citations in the paper: e.g. `&` vs `and`, and a 30-char truncation that turns
# "Urological Interests" into "Urological Interes". List multiple substrings per
# case so we match defensively; any one hit counts as recovery.
NAMED_CASE_MATCHERS = {
    "Lindeen v. SEC": ["lindeen"],
    "Ass'n of Private Sector Colleges and Universities v. Duncan": [
        "private sector colleges",
        "apscu",
    ],
    "Council for Urological Interests v. Burwell": [
        "urological interes",  # truncated in CourtListener
    ],
    # Buffington is Federal Circuit, excluded from D.C. Circuit dataset
}


def _load_classified_with_cluster() -> list[dict]:
    """Load classified records and attach cluster_id from the source opinion file.

    Older classified records don't include cluster_id; we join it in from
    data/opinions/{id}.json so downstream code can dedupe by cluster."""
    path = DATA / "classified_opinions.jsonl"
    if not path.exists():
        return []
    records = [json.loads(l) for l in path.read_text().splitlines()]
    for rec in records:
        if rec.get("cluster_id"):
            continue
        src = DATA / "opinions" / f"{rec.get('id')}.json"
        if src.exists():
            try:
                rec["cluster_id"] = json.loads(src.read_text()).get("cluster_id")
            except Exception:
                pass
    return records


def _best_record_per_cluster(records: list[dict]) -> list[dict]:
    """One record per cluster, preferring affirmance==True and highest confidence.

    CourtListener clusters can hold majority + concurrence + dissent, each
    classified separately. For counting purposes and for Stage 2 FR history
    fetches, we want the single 'best' record per cluster — the one that
    represents the court's holding."""
    def score(r: dict) -> tuple[int, float]:
        return (
            1 if r.get("is_chevron_step_two_affirmance") else 0,
            float(r.get("confidence") or 0),
        )
    by_cluster: dict = {}
    for rec in records:
        cid = rec.get("cluster_id") or rec.get("id")
        if cid not in by_cluster or score(rec) > score(by_cluster[cid]):
            by_cluster[cid] = rec
    return list(by_cluster.values())


def validate_named_cases() -> None:
    records = _load_classified_with_cluster()
    if not records:
        print("run classify first")
        return
    deduped = _best_record_per_cluster(records)
    print(f"classified opinions: {len(records)}  (deduped to {len(deduped)} clusters)")

    print()
    print("Recovered named cases:")
    missing = []
    for display, substrings in NAMED_CASE_MATCHERS.items():
        matched = None
        for rec in deduped:
            name = (rec.get("case_name") or "").lower()
            if any(s in name for s in substrings):
                matched = rec
                break
        if matched:
            print(f"  {display}")
            print(f"    stored_name: {matched.get('case_name')}")
            print(f"    step2_aff={matched.get('is_chevron_step_two_affirmance')} "
                  f"nc_rule={matched.get('is_notice_and_comment_rule')} "
                  f"conf={matched.get('confidence')}")
            print(f"    rule: {matched.get('agency')} / {matched.get('rule_short_name')}")
        else:
            missing.append(display)
    if missing:
        print()
        print("MISSING:", missing, "-- iterate on the classifier prompt")
    else:
        print()
        print("ALL NAMED CASES RECOVERED.")

    affirmed = [
        r for r in deduped
        if r.get("is_chevron_step_two_affirmance")
        and r.get("is_notice_and_comment_rule")
        and float(r.get("confidence") or 0) >= 0.6
    ]
    print()
    print(f"affirmed step-two N&C rules (deduped by cluster, conf>=0.6): {len(affirmed)}")
    for r in sorted(affirmed, key=lambda x: x.get("date_filed") or ""):
        print(f"  {r.get('date_filed')}  conf={float(r.get('confidence') or 0):.2f}  {r.get('case_name')}")
        print(f"    -> {r.get('agency')} / {r.get('rule_short_name')}")


# ------------------------------------------------------------------ stage 2a: rule histories

def pull_rule_histories() -> None:
    """For each affirmed rule, pull all subsequent Federal Register documents touching the same CFR parts.

    Deduped by cluster so a majority + concurrence + dissent triple doesn't
    trigger three identical FR-history fetches."""
    deduped = _best_record_per_cluster(_load_classified_with_cluster())
    for rec in deduped:
        if not rec.get("is_chevron_step_two_affirmance") or float(rec.get("confidence") or 0) < 0.6:
            continue
        if not rec.get("is_notice_and_comment_rule"):
            continue
        parts = extract_cfr_parts(rec.get("cfr_citation", ""))
        origin_date = rec.get("date_filed")
        if not parts:
            continue
        for part in parts:
            fetch_fr_by_cfr(part, origin_date, rec["id"])


def extract_cfr_parts(cite: str) -> list[str]:
    """Turn a CFR citation string into ['TITLE CFR PART', ...], deduped.

    Handles the common shapes we see in D.C. Circuit opinions and LLM output:
        '34 C.F.R. pts. 600, 668'     -> ['34 CFR 600', '34 CFR 668']
        '34 CFR Part 600'             -> ['34 CFR 600']
        '34 C.F.R. parts 600 and 668' -> ['34 CFR 600', '34 CFR 668']
        '34 C.F.R. \u00a7\u00a7 668.7, 668.204' -> ['34 CFR 668']
        '34 C.F.R. \u00a7 668.7(a)'   -> ['34 CFR 668']
    Non-CFR strings (e.g. Federal Register cites) return [].
    """
    if not cite:
        return []
    m = re.search(r"(\d+)\s*C\.?\s*F\.?\s*R\.?", cite, re.IGNORECASE)
    if not m:
        return []
    title = m.group(1)
    tail = cite[m.end():]
    parts: list[str] = []
    seen: set[str] = set()
    # Match a part number (1-4 digits) optionally followed by .subsection; we
    # only keep the part, so '668.7' and '668.204' both yield part '668'.
    for token in re.findall(r"\b(\d{1,4})(?:\.\d+)?\b", tail):
        if token in seen:
            continue
        seen.add(token)
        parts.append(f"{title} CFR {token}")
    return parts


def fetch_fr_by_cfr(cfr_spec: str, after_date: str, origin_opinion_id: int) -> None:
    """Query the Federal Register API for all documents affecting this CFR part after after_date."""
    title, _, part = cfr_spec.split(" ")
    url = f"{FEDERAL_REGISTER_BASE}/documents.json"
    params = {
        "conditions[cfr][title]": title,
        "conditions[cfr][part]": part,
        "conditions[publication_date][gte]": after_date,
        "order": "oldest",
        "per_page": 100,
    }
    while url:
        r = requests.get(url, params=params, timeout=30)
        r.raise_for_status()
        data = r.json()
        for doc in data.get("results", []):
            doc_number = doc.get("document_number")
            if not doc_number:
                continue
            dst = DATA / "amendments" / f"{origin_opinion_id}__{doc_number}.json"
            if dst.exists():
                continue
            # augment with abstract + full_text_xml_url for the classifier
            dst.write_text(json.dumps({"origin_opinion_id": origin_opinion_id, **doc}, indent=2))
        url = data.get("next_page_url")
        params = None
        time.sleep(0.3)


# ------------------------------------------------------------------ stage 2b: classify amendments

AMENDMENT_CLASSIFIER_FALLBACK = """You are coding a subsequent Federal Register amendment against a prior agency rule
that was affirmed by the D.C. Circuit under Chevron step two.

Assign EXACTLY ONE of these nine categories (from Bressman & Stack 2025):

  no_revisions              - not applicable here; skip
  technical_correction      - minor, non-substantive (e.g., fixing a citation, renaming an exchange)
  revised_date              - changes an effective or compliance date; substance unchanged
  revised_paperwork         - changes record-keeping / reporting / disclosure mechanics
  clarification             - clarifies an existing interpretation without changing it
  unrelated_amendment       - touches the rule's CFR part but not the interpretation the court affirmed
  new_application           - applies the affirmed interpretation to new facts or a new subject
  additional_factor         - adds a consideration not inconsistent with the affirmed interpretation
  reversal                  - adopts a wholly inconsistent interpretation of the statute

Return JSON:

  category: one of the above
  is_reversal: boolean
  justification: short passage quoted from the amendment text supporting your choice
  confidence: 0.0 to 1.0

AFFIRMED RULE (from the D.C. Circuit decision):
{{AFFIRMED_RULE_SUMMARY}}

SUBSEQUENT AMENDMENT ABSTRACT:
{{AMENDMENT_ABSTRACT}}

SUBSEQUENT AMENDMENT FULL TEXT (first 20,000 chars):
{{AMENDMENT_TEXT}}

Respond with ONLY the JSON object.
"""


def classify_amendments() -> None:
    """Classify every pulled amendment via the Batches API."""
    import anthropic
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    static_prefix, varying_template = _split_for_cache(AMENDMENT_CLASSIFIER_FALLBACK)

    # load affirmed-rule summaries so each amendment request carries the
    # context of the rule it's being compared against.
    rule_summaries: dict = {}
    for line in (DATA / "classified_opinions.jsonl").read_text().splitlines():
        rec = json.loads(line)
        rule_summaries[rec["id"]] = (
            f"{rec.get('case_name')} \u2014 {rec.get('agency')} {rec.get('rule_short_name')}. "
            f"Court upheld at Chevron step two. Reasoning: {rec.get('reasoning','')}"
        )

    out_path = DATA / "classified_amendments.jsonl"
    seen: set = set()
    if out_path.exists():
        for line in out_path.read_text().splitlines():
            seen.add(json.loads(line)["amendment_key"])

    batch_requests: list = []
    key_to_amend: dict = {}
    for amend_file in sorted((DATA / "amendments").glob("*.json")):
        key = amend_file.stem
        if key in seen:
            continue
        data = json.loads(amend_file.read_text())
        origin = data.get("origin_opinion_id")
        full_text = fetch_fr_full_text(data.get("full_text_xml_url")) or data.get("body") or ""
        varying = (
            varying_template
            .replace("{{AFFIRMED_RULE_SUMMARY}}", rule_summaries.get(origin, "(unknown)"))
            .replace("{{AMENDMENT_ABSTRACT}}", data.get("abstract") or "")
            .replace("{{AMENDMENT_TEXT}}", full_text[:20_000])
        )
        # Batch custom_id has a 64-char limit; amendment_key is usually short
        # but clip defensively and encode to stay within [A-Za-z0-9_-].
        custom_id = ("am-" + re.sub(r"[^A-Za-z0-9_-]", "_", key))[:64]
        key_to_amend[custom_id] = (key, data, origin)
        batch_requests.append({
            "custom_id": custom_id,
            "params": {
                "model": MODEL,
                "max_tokens": 500,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "text", "text": static_prefix, "cache_control": {"type": "ephemeral"}},
                        {"type": "text", "text": varying},
                    ],
                }],
            },
        })

    if not batch_requests:
        print("no new amendments to classify")
        return
    print(f"queueing {len(batch_requests)} amendments for classification via Batches API")

    with out_path.open("a") as out:
        for i in range(0, len(batch_requests), _BATCH_CHUNK):
            chunk = batch_requests[i:i + _BATCH_CHUNK]
            label = f"amendments {i+1}-{i+len(chunk)}"
            batch = _submit_and_wait_for_batch(client, chunk, label)
            for result in client.messages.batches.results(batch.id):
                entry = key_to_amend.get(result.custom_id)
                if not entry:
                    continue
                key, data, origin = entry
                if result.result.type != "succeeded":
                    err = getattr(getattr(result.result, "error", None), "type", result.result.type)
                    print(f"  {result.custom_id} failed: {err}")
                    continue
                msg = result.result.message
                parsed = parse_json_block(msg.content[0].text)
                record = {
                    "amendment_key": key,
                    "origin_opinion_id": origin,
                    "document_number": data.get("document_number"),
                    "publication_date": data.get("publication_date"),
                    "agency_names": data.get("agency_names"),
                    "title": data.get("title"),
                    **parsed,
                }
                out.write(json.dumps(record) + "\n")
                out.flush()


def fetch_fr_full_text(url: str | None) -> str:
    if not url:
        return ""
    try:
        r = requests.get(url, timeout=30)
        r.raise_for_status()
        # crude XML -> text; replace with lxml for production
        return re.sub(r"<[^>]+>", " ", r.text)
    except Exception:
        return ""


# ------------------------------------------------------------------ stage 3: package

def package() -> None:
    import pandas as pd  # pip install pandas openpyxl
    import matplotlib.pyplot as plt  # pip install matplotlib

    opinions = _best_record_per_cluster(_load_classified_with_cluster())
    amendments = [json.loads(l) for l in (DATA / "classified_amendments.jsonl").read_text().splitlines()]

    affirmed = [
        o for o in opinions
        if o.get("is_chevron_step_two_affirmance")
        and o.get("is_notice_and_comment_rule")
        and float(o.get("confidence") or 0) >= 0.6
    ]
    pd.DataFrame(affirmed).to_excel(OUT / "cases.xlsx", index=False)
    pd.DataFrame(amendments).to_excel(OUT / "amendments.xlsx", index=False)

    # chart: amendment counts by category
    cats = pd.Series([a.get("category", "unknown") for a in amendments]).value_counts()
    fig, ax = plt.subplots(figsize=(9, 5))
    colors = ["#c00" if c == "reversal" else "#888" for c in cats.index]
    ax.bar(cats.index, cats.values, color=colors)
    ax.set_ylabel("count")
    ax.set_title("Subsequent amendments to rules affirmed under Chevron step two (D.C. Cir. 2004–2024)")
    plt.xticks(rotation=30, ha="right")
    plt.tight_layout()
    plt.savefig(OUT / "chart.png", dpi=150)

    write_memo(affirmed, amendments)


def write_memo(affirmed, amendments):
    from docx import Document  # pip install python-docx
    n_cases = len(affirmed)
    reversals = [a for a in amendments if a.get("is_reversal")]
    d = Document()
    d.add_heading("Computational replication of Bressman & Stack (2025)", level=1)
    d.add_paragraph(
        f"Using free public sources (CourtListener, Federal Register API) and an LLM classifier, "
        f"we identified {n_cases} D.C. Circuit decisions (2004–2024) upholding an agency's "
        f"interpretation of a notice-and-comment rule under Chevron step two. The authors found 96. "
        f"Our count differs for the reasons documented in BRIEFING.md §5."
    )
    d.add_paragraph(
        f"Of the subsequent Federal Register amendments to those rules, {len(reversals)} were coded "
        f"as reversals by the classifier. The authors found 3. The chart in chart.png shows the "
        f"full category distribution."
    )
    d.add_heading("Method", level=2)
    d.add_paragraph(
        "See BRIEFING.md for the full methodology. Two automated stages: Stage 1 searches CourtListener "
        "for D.C. Circuit opinions mentioning Chevron, then uses Claude to classify each as a step-two "
        "affirmance of a notice-and-comment rule. Stage 2 pulls each rule's Federal Register history "
        "via CFR-part queries and has Claude code each subsequent document into one of the nine "
        "Bressman-Stack categories."
    )
    d.add_heading("Limitations", level=2)
    d.add_paragraph(
        "LLM classification is probabilistic; expect disagreement with an expert coder on 5–15% of "
        "borderline cases. CourtListener coverage is not identical to Westlaw. CFR-part matching "
        "over-collects amendments; the classifier is expected to code most as 'unrelated'. This is "
        "an illustrative replication, not a publishable empirical study."
    )
    d.save(OUT / "memo.docx")


# ------------------------------------------------------------------ CLI

def main():
    if len(sys.argv) < 2:
        print(__doc__)
        return
    cmd = sys.argv[1]
    if cmd == "pull":
        pull_opinions()
    elif cmd == "classify":
        classify_opinions()
    elif cmd == "validate":
        validate_named_cases()
    elif cmd == "history":
        pull_rule_histories()
    elif cmd == "code-amendments":
        classify_amendments()
    elif cmd == "package":
        package()
    elif cmd == "smoke":
        pull_opinions(SMOKE_FROM, SMOKE_TO)
        classify_opinions()
        validate_named_cases()
    else:
        print(f"unknown command: {cmd}")


if __name__ == "__main__":
    main()
